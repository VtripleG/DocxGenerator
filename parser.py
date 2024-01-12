import pprint
from docx.enum.text import WD_UNDERLINE
import xmltodict
import json
from docx.document import Document

try:
    document = Document()
except TypeError:
    from docx import Document

    document = Document()


def SearchParagraph(paragraphs, str):
    count = 0
    for paragraph in paragraphs:
        count += 1
        if str in paragraph.text:
            print(count)


def __DeleteLastColumn(table_index, document):
    table = document.tables[table_index]
    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for cell in table.column_cells(-1):
        cell._tc.getparent().remove(cell._tc)
    col_elem = grid[-1]
    grid.remove(col_elem)


def __DeleteRow(item):
    item._element.getparent().remove(item._element)


def __DeleteTable(table):
    table._element.getparent().remove(table._element)


def __DeleteParagraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def XmlToDict(fileName):
    with (open(fileName, 'r', encoding='utf16') as xml_file):
        xml_data = xmltodict.parse(xml_file.read())
    return xml_data


def GenerateDocxOch(dictFullInfOchnoe: dict, doc: Document) -> Document:
    dictTimeOchnoe = dictFullInfOchnoe['Часы']

    kursRabFlagOO = False
    kursPrFlagOO = False
    konRabFlagOO = False
    practicalTrainingFlagOchnoe = False

    listKursRabOO = list()
    listKursPrOO = list()

    zedOO = int(0)

    for key in dictTimeOchnoe.keys():
        semestrDict = dictTimeOchnoe[key]
        if 'Курсовая работа' in semestrDict.keys():
            kursRabFlagOO = True
            listKursRabOO.append(key)
        if 'Курсовой проект' in semestrDict.keys():
            kursPrFlagOO = True
            listKursPrOO.append(key)
        if 'Контрольная работа' in semestrDict.keys():
            konRabFlagOO = True
        if 'ЗЕТ' in semestrDict.keys():
            zedOO += int(semestrDict['ЗЕТ'])

    stringKursRabOO = str()
    if len(listKursRabOO) != 0:
        stringKursRabOO = f"{listKursRabOO[0]} семестре для очной формы обучения"
    if len(listKursRabOO) > 1:
        for index in range(1, len(listKursRabOO)):
            stringKursRabOO += f", в {listKursRabOO[index]} семестре для очной формы обучения"

    stringKursPrOO = str()
    if len(listKursPrOO) != 0:
        stringKursPrOO = f"{listKursPrOO[0]} семестре для очной формы обучения"
    if len(listKursPrOO) > 1:
        for index in range(1, len(listKursPrOO)):
            stringKursPrOO += f", в {listKursPrOO[index]} семестре для очной формы обучения"

    dictTypeWork = dict()
    for semesterNumber in dictTimeOchnoe.keys():
        for key in dictTimeOchnoe[semesterNumber].keys():
            dictTypeWork[key] = 0
    stringTypeWork = str()
    for key in dictTypeWork.keys():
        if key == 'Практические занятия':
            stringTypeWork += ' проводятся практические занятия,'
        elif key == 'Лабораторные занятия':
            stringTypeWork += ' проводятся лабораторные работы,'
        elif key == 'Лекционные занятия':
            stringTypeWork += ' читаются лекции,'
        elif key == 'Курсовой проект':
            stringTypeWork += ' выполняется курсовой проект,'
        elif key == 'Курсовая работа':
            stringTypeWork += ' выполняется курсовая работа,'
        elif key == 'Контрольная работа':
            stringTypeWork += ' проводится контрольная работа,'
    if stringTypeWork != '':
        stringTypeWork = stringTypeWork[:-1]

    listSemestersOchnoe = dictTimeOchnoe.keys()
    stringAttistatList = str()

    for semesterNumber in listSemestersOchnoe:
        stringAttistatList += f"{semesterNumber} семестре для очной формы обучения, "
    stringAttistatList = stringAttistatList[:-2]

    for object in doc.paragraphs:
        if 'name' in object.text:
            for run in object.runs:
                run.text = run.text.replace('name', dictFullInfOchnoe['Название'])
        if 'spec' in object.text:
            for run in object.runs:
                run.text = run.text.replace('spec', dictFullInfOchnoe['Специальность'])
        if 'prof' in object.text:
            for run in object.runs:
                run.text = run.text.replace('prof', dictFullInfOchnoe['Профиль'])
        if 'qual' in object.text:
            for run in object.runs:
                run.text = run.text.replace('qual', dictFullInfOchnoe['Квалификация'])
        if 'period' in object.text:
            for run in object.runs:
                run.text = run.text.replace('period', dictFullInfOchnoe['srok'])
        if 'form' in object.text:
            for run in object.runs:
                run.text = run.text.replace('form', 'Очная')
        if 'startYear' in object.text:
            for run in object.runs:
                run.text = run.text.replace('startYear', dictFullInfOchnoe['startYear'])
        if 'B1O' in object.text:
            for run in object.runs:
                if 'B1O' in run.text:
                    if dictFullInfOchnoe['B1'][3] == 'О':
                        run.text = run.text.replace('B1O', '')
                    else:
                        run.clear()
        if 'B1B' in object.text:
            for run in object.runs:
                if 'B1B' in run.text:
                    if dictFullInfOchnoe['B1'][3] != 'О':
                        run.text = run.text.replace('B1B)', '')
                    else:
                        run.clear()
        if 'zed' in object.text:
            for run in object.runs:
                run.text = run.text.replace('zed', str(zedOO))
        if 'compList' in object.text:
            compStr = str()
            for key in dictFullInfOchnoe['Компетенции'].keys():
                compStr += key + ' - ' + dictFullInfOchnoe['Компетенции'][key] + '\n'
            object.text = object.text.replace('compList', compStr)
        if 'kursPList' in object.text:
            for run in object.runs:
                run.text = run.text.replace('kursPList', stringKursPrOO)
        if 'kursRList' in object.text:
            for run in object.runs:
                run.text = run.text.replace('kursRList', stringKursRabOO)
        if 'LabRList' in object.text:
            labRFlag = False
            for semestr in dictTimeOchnoe.keys():
                if 'Лабораторные занятия' in dictTimeOchnoe[semestr].keys():
                    labRFlag = True
            if labRFlag == True:
                __DeleteParagraph(object)
            else:
                object.text = object.text.replace('LabRList', '')
        if 'PractPodgotov' in object.text:
            if practicalTrainingFlagOchnoe == True:
                for run in object.runs:
                    run.text = run.text.replace('PractPodgotov', '')
            else:
                __DeleteParagraph(object)
        if 'KPrY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KPrY', '')
            if (kursPrFlagOO == False):
                __DeleteParagraph(object)
        if 'KRabY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KRabY', '')
            if (kursRabFlagOO == False):
                __DeleteParagraph(object)
        if 'KPY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KPY', '')
            if (kursPrFlagOO == False and kursRabFlagOO == False):
                __DeleteParagraph(object)
        if 'KPN' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KPN', '')
            if (kursPrFlagOO == True or kursRabFlagOO == True):
                __DeleteParagraph(object)
        if 'KRY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KRY', '')
            if (konRabFlagOO == False):
                __DeleteParagraph(object)
        if 'KRN' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KRN', '')
            if (konRabFlagOO == True):
                __DeleteParagraph(object)
        if 'attistList' in object.text:
            for run in object.runs:
                run.text = run.text.replace('attistList', stringAttistatList)
        if 'stringTypeWork' in object.text:
            for run in object.runs:
                run.text = run.text.replace('stringTypeWork', stringTypeWork)
        if 'timeTableZ' in object.text:
            __DeleteParagraph(object)
        if 'allTimeTableZ' in object.text:
            __DeleteParagraph(object)

    for _ in range(4 - len(dictFullInfOchnoe['Компетенции'])):
        for _ in range(3):
            __DeleteRow(doc.tables[1].rows[-1])

    compiKeysOO = list()
    for key in dictFullInfOchnoe['Компетенции'].keys():
        compiKeysOO.append(key)
    for cell in doc.tables[1].columns[0].cells:
        if 'comp' in cell.text:
            cell.text = cell.text.replace('comp', compiKeysOO[0])
            compiKeysOO.pop(0)

    for semestrKey in dictTimeOchnoe.keys():
        dictTimeOchnoe[semestrKey]['Аудиторные занятия'] = str('0')
        for timeKey in dictTimeOchnoe[semestrKey].keys():
            if timeKey == 'Практические занятия' or timeKey == 'Лабораторные занятия' or timeKey == 'Лекционные занятия':
                dictTimeOchnoe[semestrKey]['Аудиторные занятия'] = str(
                    int(dictTimeOchnoe[semestrKey]['Аудиторные занятия']) + int(dictTimeOchnoe[semestrKey][timeKey]))

    timeKeysOO = list()
    for key in dictTimeOchnoe.keys():
        timeKeysOO.append(key)

    dictAllTimeOO = dict()
    dictAllTimeOO['Аудиторные занятия'] = str(0)
    dictAllTimeOO['allTime'] = str(0)
    for keySemestr in dictTimeOchnoe.keys():
        for timeKey in dictTimeOchnoe[keySemestr].keys():
            if timeKey == 'Практические занятия' or timeKey == 'Лабораторные занятия' or timeKey == 'Лекционные занятия':
                dictAllTimeOO['Аудиторные занятия'] = str(
                    int(dictAllTimeOO['Аудиторные занятия']) + int(dictTimeOchnoe[keySemestr][timeKey]))
                try:
                    dictAllTimeOO[timeKey] = str(int(dictAllTimeOO[timeKey]) + int(dictTimeOchnoe[keySemestr][timeKey]))
                except:
                    dictAllTimeOO[timeKey] = dictTimeOchnoe[keySemestr][timeKey]
            if timeKey == 'Самостоятельная работа' or timeKey == 'Итого часов':
                try:
                    dictAllTimeOO[timeKey] = str(int(dictAllTimeOO[timeKey]) + int(dictTimeOchnoe[keySemestr][timeKey]))
                except:
                    dictAllTimeOO[timeKey] = dictTimeOchnoe[keySemestr][timeKey]
    if 'Самостоятельная работа' in dictAllTimeOO.keys():
        dictAllTimeOO['allTime'] = dictAllTimeOO['Аудиторные занятия'] + dictAllTimeOO['Самостоятельная работа']
    else:
        dictAllTimeOO['allTime'] = dictAllTimeOO['Аудиторные занятия']

    timeTableOO = doc.tables[2]
    match len(dictTimeOchnoe):
        case 1:
            for _ in range(3 * 13):
                __DeleteRow(timeTableOO.rows[-1])
        case 2:
            for _ in range(13):
                __DeleteRow(timeTableOO.rows[0])
            for _ in range(2 * 13):
                __DeleteRow(timeTableOO.rows[-1])
        case 3:
            for _ in range(2 * 13):
                __DeleteRow(timeTableOO.rows[0])
            for _ in range(13):
                __DeleteRow(timeTableOO.rows[-1])
        case 4:
            for _ in range(3 * 13):
                __DeleteRow(timeTableOO.rows[0])

    for colum in timeTableOO.columns:
        for cell in colum.cells:
            if 'adTimeAll' in cell.text:
                if 'Аудиторные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('adTimeAll', dictAllTimeOO['Аудиторные занятия'])
                else:
                    cell.text = ''
            if 'lcTimeAll' in cell.text:
                if 'Лекционные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('lcTimeAll', dictAllTimeOO['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'prctTimeAll' in cell.text:
                if 'Практические занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('prctTimeAll', dictAllTimeOO['Практические занятия'])
                else:
                    cell.text = ''
            if 'lbTimeAll' in cell.text:
                if 'Лабораторные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('lbTimeAll', dictAllTimeOO['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'smTimeAll' in cell.text:
                if 'Самостоятельная работа' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('smTimeAll', dictAllTimeOO['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'allTime' in cell.text:
                if 'Итого часов' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allTime', dictAllTimeOO['Итого часов'])
                else:
                    cell.text = ''
            if 'allZed' in cell.text:
                cell.text = str(zedOO)
            if 'semestr' in cell.text:
                cell.text = cell.text.replace('semestr', str(timeKeysOO[0]))
            if 'audTime' in cell.text:
                if 'Аудиторные занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('audTime', dictTimeOchnoe[timeKeysOO[0]]['Аудиторные занятия'])
                else:
                    cell.text = ''
            if 'practTime' in cell.text:
                if 'Практические занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('practTime', dictTimeOchnoe[timeKeysOO[0]]['Практические занятия'])
                else:
                    cell.text = ''
            if 'labTime' in cell.text:
                if 'Лабораторные занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('labTime', dictTimeOchnoe[timeKeysOO[0]]['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'lectTime' in cell.text:
                if 'Лекционные занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('lectTime', dictTimeOchnoe[timeKeysOO[0]]['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'samTime' in cell.text:
                if 'Самостоятельная работа' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('samTime', dictTimeOchnoe[timeKeysOO[0]]['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'kurs' in cell.text:
                if 'Курсовой проект' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = '+'
                else:
                    cell.text = '-'
            if 'kr' in cell.text:
                if 'Контрольная работа' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = '+'
                else:
                    cell.text = '-'
            if 'att' in cell.text:
                if 'Зачет' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('att', 'Зачет')
                elif 'Зачет с оценкой' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('att', 'Зачет с оценкой')
                elif 'Экзамен' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('att', 'Экзамен')
            if 'fullTime' in cell.text:
                if 'Итого часов' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('fullTime', dictTimeOchnoe[timeKeysOO[0]]['Итого часов'])
                else:
                    cell.text = ''
            if 'fullZed' in cell.text:
                if 'ЗЕТ' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('fullZed', dictTimeOchnoe[timeKeysOO[0]]['ЗЕТ'])
                else:
                    cell.text = ''
                timeKeysOO.pop(0)

    allTimeTableOO = doc.tables[4]
    for colum in allTimeTableOO.columns:
        for cell in colum.cells:
            if 'allLect' in cell.text:
                if 'Лекционные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allLect', dictAllTimeOO['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'allPract' in cell.text:
                if 'Практические занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allPract', dictAllTimeOO['Практические занятия'])
                else:
                    cell.text = ''
            if 'allLab' in cell.text:
                if 'Лабораторные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allLab', dictAllTimeOO['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'allSam' in cell.text:
                if 'Самостоятельная работа' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allSam', dictAllTimeOO['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'allTime' in cell.text:
                if 'Аудиторные занятия' in dictAllTimeOO.keys():
                    if 'Самостоятельная работа' in dictAllTimeOO.keys():
                        cell.text = cell.text.replace('allTime', str(int(dictAllTimeOO['Аудиторные занятия']) + int(
                            dictAllTimeOO['Самостоятельная работа'])))
                    else:
                        cell.text = ''

    for _ in range(4 - len(dictFullInfOchnoe['Компетенции'])):
        for _ in range(3):
            __DeleteRow(doc.tables[7].rows[-1])

    compiKeysOO = list()
    for key in dictFullInfOchnoe['Компетенции'].keys():
        compiKeysOO.append(key)
    for cell in doc.tables[7].columns[0].cells:
        if 'comp' in cell.text:
            cell.text = cell.text.replace('comp', compiKeysOO[0])
            compiKeysOO.pop(0)

    for _ in range(4 - len(dictFullInfOchnoe['Компетенции'])):
        for _ in range(3):
            __DeleteRow(doc.tables[8].rows[-1])

    compiKeysOO = list()
    for key in dictFullInfOchnoe['Компетенции'].keys():
        compiKeysOO.append(key)
    for cell in doc.tables[8].columns[0].cells:
        if 'comp' in cell.text:
            cell.text = cell.text.replace('comp', compiKeysOO[0])
            compiKeysOO.pop(0)

    tableCollumnType = doc.tables[10].columns[0]
    for cellNumber in range(len(tableCollumnType.cells)-1, 0, -1):
        if 'lectType' in tableCollumnType.cells[cellNumber].text:
            if 'Лекционные занятия' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('lectType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])
        elif 'practType' in tableCollumnType.cells[cellNumber].text:
            if 'Практические занятия' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('practType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])
        elif 'labType' in tableCollumnType.cells[cellNumber].text:
            if 'Лабораторные занятия' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('labType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])
        elif 'samType' in tableCollumnType.cells[cellNumber].text:
            if 'Самостоятельная работа' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('samType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])

    if practicalTrainingFlagOchnoe == False:
        __DeleteTable(doc.tables[6])

    __DeleteTable(doc.tables[5])
    __DeleteTable(doc.tables[3])

    return doc


def GenerateDocxOchZ(dictFullInfOchnoe: dict, dictFullInfZaochnoe: dict, doc: Document) -> Document:
    dictTimeOchnoe = dictFullInfOchnoe['Часы']
    dictTimeZaochnoe = dictFullInfZaochnoe['Часы']

    kursRabFlagOO = False
    kursPrFlagOO = False
    konRabFlagOO = False
    practicalTrainingFlagOchnoe = False

    kursRabFlagZO = False
    kursPrFlagZO = False
    konRabFlagZO = False
    practicalTrainingFlagZaochnoe = False

    listKursRabOO = list()
    listKursRabZO = list()
    listKursPrOO = list()
    listKursPrZO = list()

    zedOO = int(0)
    zedZO = int(0)

    for key in dictTimeOchnoe.keys():
        semestrDict = dictTimeOchnoe[key]
        if 'Курсовая работа' in semestrDict.keys():
            kursRabFlagOO = True
            listKursRabOO.append(key)
        if 'Курсовой проект' in semestrDict.keys():
            kursPrFlagOO = True
            listKursPrOO.append(key)
        if 'Контрольная работа' in semestrDict.keys():
            konRabFlagOO = True
        if 'ЗЕТ' in semestrDict.keys():
            zedOO += int(semestrDict['ЗЕТ'])
        if '' in semestrDict.keys():
            practicalTrainingFlagOchnoe = True

    for key in dictTimeZaochnoe.keys():
        semestrDict = dictTimeZaochnoe[key]
        if 'Курсовая работа' in semestrDict.keys():
            kursRabFlagZO = True
            listKursRabOO.append(key)
        if 'Курсовой проект' in semestrDict.keys():
            kursPrFlagZO = True
            listKursPrZO.append(key)
        if 'Контрольная работа' in semestrDict.keys():
            konRabFlagZO = True
        if 'ЗЕТ' in semestrDict.keys():
            zedZO += int(semestrDict['ЗЕТ'])
        if '' in semestrDict.keys():
            practicalTrainingFlagZaochnoe = True

    stringKursRabOO = str()
    if len(listKursRabOO) != 0:
        stringKursRabOO = f"{listKursRabOO[0]} семестре для очной формы обучения"
    if len(listKursRabOO) > 1:
        for index in range(1, len(listKursRabOO)):
            stringKursRabOO += f", в {listKursRabOO[index]} семестре для очной формы обучения"

    stringKursRabZO = str()
    if len(listKursRabZO) != 0:
        if stringKursRabZO != '':
            for index in range(0, len(listKursRabZO)):
                stringKursRabZO += f", в {listKursRabZO[index]} семестре для заочной формы обучения"
        else:
            stringKursRabZO = f"{listKursRabZO[0]} семестре для заочной формы обучения"
            if len(listKursRabZO) > 1:
                for index in range(1, len(listKursRabZO)):
                    stringKursRabZO += f", в {listKursRabZO[index]} семестре для заочной формы обучения"

    stringKursRab = stringKursRabOO + stringKursRabZO

    stringKursPrOO = str()
    if len(listKursPrOO) != 0:
        stringKursPrOO = f"{listKursPrOO[0]} семестре для очной формы обучения"
    if len(listKursPrOO) > 1:
        for index in range(1, len(listKursPrOO)):
            stringKursPrOO += f", в {listKursPrOO[index]} семестре для очной формы обучения"
    stringKursPrZO = str()
    if len(listKursPrZO) != 0:
        if stringKursPrOO != '':
            for index in range(0, len(listKursPrZO)):
                stringKursPrZO += f", в {listKursPrZO[index]} семестре для заочной формы обучения"
        else:
            stringKursPrZO = f"{listKursPrZO[0]} семестре для заочной формы обучения"
            if len(listKursPrZO) > 1:
                for index in range(1, len(listKursPrZO)):
                    stringKursPrZO += f", в {listKursPrZO[index]} семестре для заочной формы обучения"

    stringKursPr = stringKursPrOO + stringKursPrZO

    dictTypeWork = dict()
    for semesterNumber in dictTimeOchnoe.keys():
        for key in dictTimeOchnoe[semesterNumber].keys():
            dictTypeWork[key] = 0
    stringTypeWork = str()
    for key in dictTypeWork.keys():
        if key == 'Практические занятия':
            stringTypeWork+= ' проводятся практические занятия,'
        elif key == 'Лабораторные занятия':
            stringTypeWork+= ' проводятся лабораторные работы,'
        elif key == 'Лекционные занятия':
            stringTypeWork+= ' читаются лекции,'
        elif key == 'Курсовой проект':
            stringTypeWork+= ' выполняется курсовой проект,'
        elif key == 'Курсовая работа':
            stringTypeWork+= ' выполняется курсовая работа,'
        elif key == 'Контрольная работа':
            stringTypeWork+= ' проводится контрольная работа,'
    if stringTypeWork != '':
        stringTypeWork = stringTypeWork[:-1]

    listSemestersOchnoe = dictTimeOchnoe.keys()
    listSemestersZaochnoe = dictTimeZaochnoe.keys()
    stringAttistatList = str()

    for semesterNumber in listSemestersOchnoe:
        stringAttistatList+= f"{semesterNumber} семестре для очной формы обучения, "
    for semesterNumber in listSemestersZaochnoe:
        stringAttistatList+= f"{semesterNumber} семестре для заочной формы обучения, "
    stringAttistatList = stringAttistatList[:-2]

    for object in doc.paragraphs:
        if 'name' in object.text:
            for run in object.runs:
                run.text = run.text.replace('name', dictFullInfOchnoe['Название'])
        if 'spec' in object.text:
            for run in object.runs:
                run.text = run.text.replace('spec', dictFullInfOchnoe['Специальность'])
        if 'prof' in object.text:
            for run in object.runs:
                run.text = run.text.replace('prof', dictFullInfOchnoe['Профиль'])
        if 'qual' in object.text:
            for run in object.runs:
                run.text = run.text.replace('qual', dictFullInfOchnoe['Квалификация'])
        if 'period' in object.text:
            for run in object.runs:
                run.text = run.text.replace('period', f"{dictFullInfOchnoe['srok']} / {dictFullInfZaochnoe['srok']}")
        if 'form' in object.text:
            for run in object.runs:
                run.text = run.text.replace('form', 'Очная/заочная')
        if 'startYear' in object.text:
            for run in object.runs:
                run.text = run.text.replace('startYear', dictFullInfOchnoe['startYear'])
        if 'B1O' in object.text:
            for run in object.runs:
                if 'B1O' in run.text:
                    if dictFullInfOchnoe['B1'][3] == 'О':
                        run.text = run.text.replace('B1O', '')
                    else:
                        run.clear()
        if 'B1B' in object.text:
            for run in object.runs:
                if 'B1B' in run.text:
                    if dictFullInfOchnoe['B1'][3] != 'О':
                        run.text = run.text.replace('B1B)', '')
                    else:
                        run.clear()
        if 'zed' in object.text:
            for run in object.runs:
                run.text = run.text.replace('zed', f"{zedOO}")
        if 'compList' in object.text:
            compStr = str()
            for key in dictFullInfOchnoe['Компетенции'].keys():
                compStr += key + ' - ' + dictFullInfOchnoe['Компетенции'][key] + '\n'
            object.text = object.text.replace('compList', compStr)
        if 'kursPList' in object.text:
            for run in object.runs:
                run.text = run.text.replace('kursPList', stringKursPr)
        if 'kursRList' in object.text:
            for run in object.runs:
                run.text = run.text.replace('kursRList', stringKursRab)
        if 'LabRList' in object.text:
            labRFlag = False
            for semestr in dictTimeOchnoe.keys():
                if 'Лабораторные занятия' in dictTimeOchnoe[semestr].keys():
                    labRFlag = True
            if labRFlag == True:
                __DeleteParagraph(object)
            else:
                object.text = object.text.replace('LabRList', '')
        if 'PractPodgotov' in object.text:
            if (practicalTrainingFlagOchnoe == True) or (practicalTrainingFlagZaochnoe == True):
                for run in object.runs:
                    run.text = run.text.replace('PractPodgotov', '')
            else:
                __DeleteParagraph(object)
        if 'KPrY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KPrY', '')
            if (kursPrFlagOO == False and kursPrFlagZO == False):
                __DeleteParagraph(object)
        if 'KRabY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KRabY', '')
            if (kursRabFlagZO == False and kursRabFlagOO == False):
                __DeleteParagraph(object)
        if 'KPY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KPY', '')
            if (kursPrFlagOO == False and kursPrFlagZO == False and kursRabFlagOO == False and kursRabFlagZO == False):
                __DeleteParagraph(object)
        if 'KPN' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KPN', '')
            if (kursPrFlagOO == True or kursPrFlagZO == True or kursRabFlagOO == True):
                __DeleteParagraph(object)
        if 'KRY' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KRY', '')
            if (konRabFlagOO == False and konRabFlagZO == False):
                __DeleteParagraph(object)
        if 'KRN' in object.text:
            for run in object.runs:
                run.text = run.text.replace('KRN', '')
            if (konRabFlagOO == True or konRabFlagZO == True):
                __DeleteParagraph(object)
        if 'attistList' in object.text:
            for run in object.runs:
                run.text = run.text.replace('attistList', stringAttistatList)
        if 'stringTypeWork' in object.text:
            for run in object.runs:
                run.text = run.text.replace('stringTypeWork', stringTypeWork)
        if 'timeTableZ' in object.text:
            object.runs[0].text = object.runs[0].text.replace('timeTableZ', '')
        if 'allTimeTableZ' in object.text:
            object.runs[0].text = object.runs[0].text.replace('allTimeTableZ', '')

    for _ in range(4 - len(dictFullInfOchnoe['Компетенции'])):
        for _ in range(3):
            __DeleteRow(doc.tables[1].rows[-1])

    compiKeysOO = list()
    for key in dictFullInfOchnoe['Компетенции'].keys():
        compiKeysOO.append(key)
    for cell in doc.tables[1].columns[0].cells:
        if 'comp' in cell.text:
            cell.text = cell.text.replace('comp', compiKeysOO[0])
            compiKeysOO.pop(0)

    for semestrKey in dictTimeOchnoe.keys():
        dictTimeOchnoe[semestrKey]['Аудиторные занятия'] = str('0')
        for timeKey in dictTimeOchnoe[semestrKey].keys():
            if timeKey == 'Практические занятия' or timeKey == 'Лабораторные занятия' or timeKey == 'Лекционные занятия':
                dictTimeOchnoe[semestrKey]['Аудиторные занятия'] = str(
                    int(dictTimeOchnoe[semestrKey]['Аудиторные занятия']) + int(dictTimeOchnoe[semestrKey][timeKey]))

    timeKeysOO = list()
    for key in dictTimeOchnoe.keys():
        timeKeysOO.append(key)

    dictAllTimeOO = dict()
    dictAllTimeOO['Аудиторные занятия'] = str(0)
    dictAllTimeOO['allTime'] = str(0)
    for keySemestr in dictTimeOchnoe.keys():
        for timeKey in dictTimeOchnoe[keySemestr].keys():
            if timeKey == 'Практические занятия' or timeKey == 'Лабораторные занятия' or timeKey == 'Лекционные занятия':
                dictAllTimeOO['Аудиторные занятия'] = str(
                    int(dictAllTimeOO['Аудиторные занятия']) + int(dictTimeOchnoe[keySemestr][timeKey]))
                try:
                    dictAllTimeOO[timeKey] = str(int(dictAllTimeOO[timeKey]) + int(dictTimeOchnoe[keySemestr][timeKey]))
                except:
                    dictAllTimeOO[timeKey] = dictTimeOchnoe[keySemestr][timeKey]
            if timeKey == 'Самостоятельная работа' or timeKey == 'Итого часов':
                try:
                    dictAllTimeOO[timeKey] = str(int(dictAllTimeOO[timeKey]) + int(dictTimeOchnoe[keySemestr][timeKey]))
                except:
                    dictAllTimeOO[timeKey] = dictTimeOchnoe[keySemestr][timeKey]
    if 'Самостоятельная работа' in dictAllTimeOO.keys():
        dictAllTimeOO['allTime'] = dictAllTimeOO['Аудиторные занятия'] + dictAllTimeOO['Самостоятельная работа']
    else:
        dictAllTimeOO['allTime'] = dictAllTimeOO['Аудиторные занятия']

    timeTableOO = doc.tables[2]
    match len(dictTimeOchnoe):
        case 1:
            for _ in range(3 * 13):
                __DeleteRow(timeTableOO.rows[-1])
        case 2:
            for _ in range(13):
                __DeleteRow(timeTableOO.rows[0])
            for _ in range(2 * 13):
                __DeleteRow(timeTableOO.rows[-1])
        case 3:
            for _ in range(2 * 13):
                __DeleteRow(timeTableOO.rows[0])
            for _ in range(13):
                __DeleteRow(timeTableOO.rows[-1])
        case 4:
            for _ in range(3 * 13):
                __DeleteRow(timeTableOO.rows[0])

    for colum in timeTableOO.columns:
        for cell in colum.cells:
            if 'adTimeAll' in cell.text:
                if 'Аудиторные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('adTimeAll', dictAllTimeOO['Аудиторные занятия'])
                else:
                    cell.text = ''
            if 'lcTimeAll' in cell.text:
                if 'Лекционные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('lcTimeAll', dictAllTimeOO['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'prctTimeAll' in cell.text:
                if 'Практические занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('prctTimeAll', dictAllTimeOO['Практические занятия'])
                else:
                    cell.text = ''
            if 'lbTimeAll' in cell.text:
                if 'Лабораторные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('lbTimeAll', dictAllTimeOO['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'smTimeAll' in cell.text:
                if 'Самостоятельная работа' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('smTimeAll', dictAllTimeOO['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'allTime' in cell.text:
                if 'Итого часов' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allTime', dictAllTimeOO['Итого часов'])
                else:
                    cell.text = ''
            if 'allZed' in cell.text:
                cell.text = str(zedOO)
            if 'semestr' in cell.text:
                cell.text = cell.text.replace('semestr', str(timeKeysOO[0]))
            if 'audTime' in cell.text:
                if 'Аудиторные занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('audTime', dictTimeOchnoe[timeKeysOO[0]]['Аудиторные занятия'])
                else:
                    cell.text = ''
            if 'practTime' in cell.text:
                if 'Практические занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('practTime', dictTimeOchnoe[timeKeysOO[0]]['Практические занятия'])
                else:
                    cell.text = ''
            if 'labTime' in cell.text:
                if 'Лабораторные занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('labTime', dictTimeOchnoe[timeKeysOO[0]]['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'lectTime' in cell.text:
                if 'Лекционные занятия' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('lectTime', dictTimeOchnoe[timeKeysOO[0]]['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'samTime' in cell.text:
                if 'Самостоятельная работа' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('samTime', dictTimeOchnoe[timeKeysOO[0]]['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'kurs' in cell.text:
                if ('Курсовой проект' or 'Курсовая работа') in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = '+'
                else:
                    cell.text = '-'
            if 'kr' in cell.text:
                if 'Контрольная работа' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = '+'
                else:
                    cell.text = '-'
            if 'att' in cell.text:
                if 'Зачет' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('att', 'Зачет')
                elif 'Зачет с оценкой' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('att', 'Зачет с оценкой')
                elif 'Экзамен' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('att', 'Экзамен')
            if 'fullTime' in cell.text:
                if 'Итого часов' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('fullTime', dictTimeOchnoe[timeKeysOO[0]]['Итого часов'])
                else:
                    cell.text = ''
            if 'fullZed' in cell.text:
                if 'ЗЕТ' in dictTimeOchnoe[timeKeysOO[0]].keys():
                    cell.text = cell.text.replace('fullZed', dictTimeOchnoe[timeKeysOO[0]]['ЗЕТ'])
                else:
                    cell.text = ''
                timeKeysOO.pop(0)

    allTimeTableOO = doc.tables[4]
    for colum in allTimeTableOO.columns:
        for cell in colum.cells:
            if 'allLect' in cell.text:
                if 'Лекционные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allLect', dictAllTimeOO['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'allPract' in cell.text:
                if 'Практические занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allPract', dictAllTimeOO['Практические занятия'])
                else:
                    cell.text = ''
            if 'allLab' in cell.text:
                if 'Лабораторные занятия' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allLab', dictAllTimeOO['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'allSam' in cell.text:
                if 'Самостоятельная работа' in dictAllTimeOO.keys():
                    cell.text = cell.text.replace('allSam', dictAllTimeOO['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'allTime' in cell.text:
                if 'Аудиторные занятия' in dictAllTimeOO.keys():
                    if 'Самостоятельная работа' in dictAllTimeOO.keys():
                        cell.text = cell.text.replace('allTime', str(int(dictAllTimeOO['Аудиторные занятия']) + int(
                            dictAllTimeOO['Самостоятельная работа'])))
                    else:
                        cell.text = ''

    for semestrKey in dictTimeZaochnoe.keys():
        dictTimeZaochnoe[semestrKey]['Аудиторные занятия'] = str('0')
        for timeKey in dictTimeZaochnoe[semestrKey].keys():
            if timeKey == 'Практические занятия' or timeKey == 'Лабораторные занятия' or timeKey == 'Лекционные занятия':
                dictTimeZaochnoe[semestrKey]['Аудиторные занятия'] = str(
                    int(dictTimeZaochnoe[semestrKey]['Аудиторные занятия']) + int(dictTimeZaochnoe[semestrKey][timeKey]))

    timeKeysZO = list()
    for key in dictTimeZaochnoe.keys():
        timeKeysZO.append(key)

    dictAllTimeZO = dict()
    dictAllTimeZO['Аудиторные занятия'] = str(0)
    dictAllTimeZO['allTime'] = str(0)
    for keySemestr in dictTimeZaochnoe.keys():
        for timeKey in dictTimeZaochnoe[keySemestr].keys():
            if timeKey == 'Практические занятия' or timeKey == 'Лабораторные занятия' or timeKey == 'Лекционные занятия':
                dictAllTimeZO['Аудиторные занятия'] = str(
                    int(dictAllTimeZO['Аудиторные занятия']) + int(dictTimeZaochnoe[keySemestr][timeKey]))
                try:
                    dictAllTimeZO[timeKey] = str(int(dictAllTimeZO[timeKey]) + int(dictTimeZaochnoe[keySemestr][timeKey]))
                except:
                    dictAllTimeZO[timeKey] = dictTimeZaochnoe[keySemestr][timeKey]
            if timeKey == 'Самостоятельная работа' or timeKey == 'Итого часов':
                try:
                    dictAllTimeZO[timeKey] = str(int(dictAllTimeZO[timeKey]) + int(dictTimeZaochnoe[keySemestr][timeKey]))
                except:
                    dictAllTimeZO[timeKey] = dictTimeZaochnoe[keySemestr][timeKey]
    if 'Самостоятельная работа' in dictAllTimeZO.keys():
        dictAllTimeZO['allTime'] = dictAllTimeZO['Аудиторные занятия'] + dictAllTimeZO['Самостоятельная работа']
    else:
        dictAllTimeZO['allTime'] = dictAllTimeZO['Аудиторные занятия']

    timeTableZO = doc.tables[3]
    match len(dictTimeZaochnoe):
        case 1:
            for _ in range(3 * 13):
                __DeleteRow(timeTableZO.rows[-1])
        case 2:
            for _ in range(13):
                __DeleteRow(timeTableZO.rows[0])
            for _ in range(2 * 13):
                __DeleteRow(timeTableZO.rows[-1])
        case 3:
            for _ in range(2 * 13):
                __DeleteRow(timeTableZO.rows[0])
            for _ in range(13):
                __DeleteRow(timeTableZO.rows[-1])
        case 4:
            for _ in range(3 * 13):
                __DeleteRow(timeTableZO.rows[0])

    for colum in timeTableZO.columns:
        for cell in colum.cells:
            if 'adTimeAll' in cell.text:
                if 'Аудиторные занятия' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('adTimeAll', dictAllTimeZO['Аудиторные занятия'])
                else:
                    cell.text = ''
            if 'lcTimeAll' in cell.text:
                if 'Лекционные занятия' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('lcTimeAll', dictAllTimeZO['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'prctTimeAll' in cell.text:
                if 'Практические занятия' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('prctTimeAll', dictAllTimeZO['Практические занятия'])
                else:
                    cell.text = ''
            if 'lbTimeAll' in cell.text:
                if 'Лабораторные занятия' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('lbTimeAll', dictAllTimeZO['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'smTimeAll' in cell.text:
                if 'Самостоятельная работа' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('smTimeAll', dictAllTimeZO['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'allTime' in cell.text:
                if 'Итого часов' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('allTime', dictAllTimeZO['Итого часов'])
                else:
                    cell.text = ''
            if 'allZed' in cell.text:
                cell.text = str(zedZO)
            if 'semestr' in cell.text:
                cell.text = cell.text.replace('semestr', str(timeKeysZO[0]))
            if 'audTime' in cell.text:
                if 'Аудиторные занятия' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('audTime', dictTimeZaochnoe[timeKeysZO[0]]['Аудиторные занятия'])
                else:
                    cell.text = ''
            if 'practTime' in cell.text:
                if 'Практические занятия' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('practTime', dictTimeZaochnoe[timeKeysZO[0]]['Практические занятия'])
                else:
                    cell.text = ''
            if 'labTime' in cell.text:
                if 'Лабораторные занятия' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('labTime', dictTimeZaochnoe[timeKeysZO[0]]['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'lectTime' in cell.text:
                if 'Лекционные занятия' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('lectTime', dictTimeZaochnoe[timeKeysZO[0]]['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'samTime' in cell.text:
                if 'Самостоятельная работа' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('samTime', dictTimeZaochnoe[timeKeysZO[0]]['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'kurs' in cell.text:
                if ('Курсовой проект' or 'Курсовая работа') in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = '+'
                else:
                    cell.text = '-'
            if 'kr' in cell.text:
                if 'Контрольная работа' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = '+'
                else:
                    cell.text = '-'
            if 'att' in cell.text:
                if 'Зачет' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('att', 'Зачет')
                elif 'Зачет с оценкой' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('att', 'Зачет с оценкой')
                elif 'Экзамен' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('att', 'Экзамен')
            if 'fullTime' in cell.text:
                if 'Итого часов' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('fullTime', dictTimeZaochnoe[timeKeysZO[0]]['Итого часов'])
                else:
                    cell.text = ''
            if 'fullZed' in cell.text:
                if 'ЗЕТ' in dictTimeZaochnoe[timeKeysZO[0]].keys():
                    cell.text = cell.text.replace('fullZed', dictTimeZaochnoe[timeKeysZO[0]]['ЗЕТ'])
                else:
                    cell.text = ''
                timeKeysZO.pop(0)

    allTimeTableZO = doc.tables[5]
    for colum in allTimeTableZO.columns:
        for cell in colum.cells:
            if 'allLect' in cell.text:
                if 'Лекционные занятия' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('allLect', dictAllTimeZO['Лекционные занятия'])
                else:
                    cell.text = ''
            if 'allPract' in cell.text:
                if 'Практические занятия' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('allPract', dictAllTimeZO['Практические занятия'])
                else:
                    cell.text = ''
            if 'allLab' in cell.text:
                if 'Лабораторные занятия' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('allLab', dictAllTimeZO['Лабораторные занятия'])
                else:
                    cell.text = ''
            if 'allSam' in cell.text:
                if 'Самостоятельная работа' in dictAllTimeZO.keys():
                    cell.text = cell.text.replace('allSam', dictAllTimeZO['Самостоятельная работа'])
                else:
                    cell.text = ''
            if 'allTime' in cell.text:
                if 'Аудиторные занятия' in dictAllTimeZO.keys():
                    if 'Самостоятельная работа' in dictAllTimeZO.keys():
                        cell.text = cell.text.replace('allTime', str(int(dictAllTimeZO['Аудиторные занятия']) + int(
                            dictAllTimeZO['Самостоятельная работа'])))
                    else:
                        cell.text = ''

    for _ in range(4 - len(dictFullInfOchnoe['Компетенции'])):
        for _ in range(3):
            __DeleteRow(doc.tables[7].rows[-1])

    compiKeysOO = list()
    for key in dictFullInfOchnoe['Компетенции'].keys():
        compiKeysOO.append(key)
    for cell in doc.tables[7].columns[0].cells:
        if 'comp' in cell.text:
            cell.text = cell.text.replace('comp', compiKeysOO[0])
            compiKeysOO.pop(0)

    for _ in range(4 - len(dictFullInfOchnoe['Компетенции'])):
        for _ in range(3):
            __DeleteRow(doc.tables[8].rows[-1])

    compiKeysOO = list()
    for key in dictFullInfOchnoe['Компетенции'].keys():
        compiKeysOO.append(key)
    for cell in doc.tables[8].columns[0].cells:
        if 'comp' in cell.text:
            cell.text = cell.text.replace('comp', compiKeysOO[0])
            compiKeysOO.pop(0)

    tableCollumnType = doc.tables[10].columns[0]
    for cellNumber in range(len(tableCollumnType.cells)-1, 0, -1):
        if 'lectType' in tableCollumnType.cells[cellNumber].text:
            if 'Лекционные занятия' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('lectType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])
        elif 'practType' in tableCollumnType.cells[cellNumber].text:
            if 'Практические занятия' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('practType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])
        elif 'labType' in tableCollumnType.cells[cellNumber].text:
            if 'Лабораторные занятия' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('labType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])
        elif 'samType' in tableCollumnType.cells[cellNumber].text:
            if 'Самостоятельная работа' in dictTypeWork.keys():
                for run in tableCollumnType.cells[cellNumber].paragraphs[0].runs:
                    run.text = run.text.replace('samType', '')
            else:
                __DeleteRow(doc.tables[10].rows[cellNumber])

    if (practicalTrainingFlagOchnoe == False) and (practicalTrainingFlagZaochnoe == False):
        __DeleteTable(doc.tables[6])

    return doc


def GetDisciplineList(plxData: dict) -> dict:
    list = {}
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыСтроки']:
        if '@КодКафедры' in object.keys() and object['@КодКафедры'] == '82':
            list[object['@Код']] = object['@Дисциплина']
    return list


def __SearchCompetenciesByDisciplineCode(disciplineCode: str, plxData: dict) -> dict:
    compCodeList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыКомпетенцииДисциплины']:
        if object['@КодСтроки'] == disciplineCode:
            compCodeList.append(object['@КодКомпетенции'])
    dict = {}
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыКомпетенции']:
        if compCodeList.__contains__(object['@Код']):
            dict[object['@ШифрКомпетенции']] = object['@Наименование']
    return dict


def KeyFromVal(dict: dict, val):
    for key, value in dict.items():
        if value == val:
            return key


def ReadDocxTemplate(filePath: str) -> Document:
    doc = Document(filePath)
    return doc


def SaveDocx(doc: Document, fileName: str, path: str):
    fullPath = str(path) + str(fileName) + '.docx'
    doc.save(fullPath)


def __GetSpecialization(plxData: dict) -> str:
    str = plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ООП']['@Название'] + ' / ' + \
          plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ООП']['@Шифр']
    return str


def __GetB1(disciplineCode: str, plxData: dict) -> str:
    string = str()
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыСтроки']:
        if object['@Код'] == disciplineCode:
            string = object['@ДисциплинаКод']
    return string


def __GetProfile(plxData: dict) -> str:
    string = plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ООП']['ООП']['@Название']
    return string


def __GetQualification(plxData: dict) -> str:
    code = plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ООП']['@Квалификация']
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['Уровень_образования']:
        if object['@Код_записи'] == code:
            return object['@ВидПлана']


def __GetStartYear(plxData: dict) -> str:
    string = plxData['Документ']['diffgr:diffgram']['dsMMISDB']['Планы']['@ГодНачалаПодготовки']
    return string


def __GetSrok(plxData: dict) -> str:
    string = plxData['Документ']['diffgr:diffgram']['dsMMISDB']['Планы']['@СрокОбучения'] + ' года '
    if plxData['Документ']['diffgr:diffgram']['dsMMISDB']['Планы']['@СрокОбученияМесяцев'] != '0':
        string += plxData['Документ']['diffgr:diffgram']['dsMMISDB']['Планы']['@СрокОбученияМесяцев'] + ' месяцев'
    return string


def GetFullInf(disciplineName: str, disciplineCode: str, plxData: dict) -> dict:
    dictInf = {}
    dictInf['Название'] = disciplineName
    dictInf['Специальность'] = __GetSpecialization(plxData)
    dictInf['Профиль'] = __GetProfile(plxData)
    dictInf['Квалификация'] = __GetQualification(plxData)
    dictInf['Компетенции'] = __SearchCompetenciesByDisciplineCode(disciplineCode, plxData)
    dictInf['Часы'] = __SearchHours(disciplineCode, plxData)
    dictInf['B1'] = __GetB1(disciplineCode, plxData)
    dictInf['startYear'] = __GetStartYear(plxData)
    dictInf['srok'] = __GetSrok(plxData)

    return dictInf


def GetFullInfOchnoe(disciplineName: str, disciplineCode: str, plxData: dict) -> dict:
    dictInf = {}
    dictInf['Название'] = disciplineName
    dictInf['Специальность'] = __GetSpecialization(plxData)
    dictInf['Профиль'] = __GetProfile(plxData)
    dictInf['Квалификация'] = __GetQualification(plxData)
    dictInf['Компетенции'] = __SearchCompetenciesByDisciplineCode(disciplineCode, plxData)
    dictInf['Часы'] = __SearchHoursOcnoe(disciplineCode, plxData)
    dictInf['Практическая подготовка'] = __SearchPracticalJobsOchnoe(disciplineCode, plxData)
    if not dictInf['Практическая подготовка']:
        dictInf.pop('Практическая подготовка')
    dictInf['B1'] = __GetB1(disciplineCode, plxData)
    dictInf['startYear'] = __GetStartYear(plxData)
    dictInf['srok'] = __GetSrok(plxData)
    return dictInf


def GetFullInfZaochnoe(disciplineName: str, disciplineCode: str, plxData: dict) -> dict:
    dictInf = {}
    dictInf['Название'] = disciplineName
    dictInf['Специальность'] = __GetSpecialization(plxData)
    dictInf['Профиль'] = __GetProfile(plxData)
    dictInf['Квалификация'] = __GetQualification(plxData)
    dictInf['Компетенции'] = __SearchCompetenciesByDisciplineCode(disciplineCode, plxData)
    dictInf['Часы'] = __SearchHoursZaochnoe(disciplineCode, plxData)
    dictInf['Практическая подготовка'] = __SearchPracticalJobsZaochnoe(disciplineCode, plxData)
    if not dictInf['Практическая подготовка']:
        dictInf.pop('Практическая подготовка')
    if 'Практическая подготовка' in dictInf.keys():
        print(dictInf['Практическая подготовка'])
    dictInf['B1'] = __GetB1(disciplineCode, plxData)
    dictInf['startYear'] = __GetStartYear(plxData)
    dictInf['srok'] = __GetSrok(plxData)

    return dictInf


def __SearchHoursBySemesterNumber(semesterNumber: int, disciplineCode: str, plxData: dict) -> dict:
    codeList = []
    hoursList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if (object['@КодТипаЧасов'] != '3') and (object['@КодОбъекта'] == disciplineCode) and (
                (int(object['@Курс']) * 2 - 1 + int(object['@Семестр']) - 1 == semesterNumber) or (
                int(object['@Курс']) * 2 - 1 + ((int(object['@Сессия']) - 1) // 2) == semesterNumber)):
            if codeList.__contains__(object['@КодВидаРаботы']) == False:
                codeList.append(object['@КодВидаРаботы'])
                hoursList.append(object['@Количество'])
    nameList = []
    dict = {}
    for key in codeList:
        for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['СправочникВидыРабот']:
            if object['@Код'] == key:
                nameList.append(object['@Название'])
    # print(hoursList)
    # print(nameList)
    for i in range(nameList.__len__()):
        dict[nameList[i]] = hoursList[i]

    return dict


def __SearchHours(disciplineCode: str, plxData: dict) -> dict:
    dict = {}
    semesterNumberList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:

        if object['@КодОбъекта'] == disciplineCode:
            if object['@Семестр'] != '0':
                num = int(object['@Курс']) * 2 - 1 + int(object['@Семестр']) - 1
            else:
                num = int(object['@Курс']) * 2 - 1 + ((int(object['@Сессия']) - 1) // 2)
            if semesterNumberList.__contains__(num) == False:
                semesterNumberList.append(num)
    for i in range(semesterNumberList.__len__()):
        dict[semesterNumberList[i]] = __SearchHoursBySemesterNumber(semesterNumberList[i], disciplineCode, plxData)
    return dict


def __SearchHoursBySemesterNumberOchnoe(semesterNumber: int, disciplineCode: str, plxData: dict) -> dict:
    codeList = []
    hoursList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if (object['@КодТипаЧасов'] != '3' and object['@КодТипаЧасов'] != '5') and (object['@КодОбъекта'] == disciplineCode) and (
                int(object['@Курс']) * 2 - 1 + int(object['@Семестр']) - 1 == semesterNumber):
            if codeList.__contains__(object['@КодВидаРаботы']) == False:
                codeList.append(object['@КодВидаРаботы'])
                hoursList.append(object['@Количество'])
    nameList = []
    dict = {}
    for key in codeList:
        for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['СправочникВидыРабот']:
            if object['@Код'] == key:
                nameList.append(object['@Название'])
    for i in range(nameList.__len__()):
        dict[nameList[i]] = hoursList[i]

    return dict


def __SearchHoursOcnoe(disciplineCode: str, plxData: dict) -> dict:
    dict = {}
    semesterNumberList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if object['@КодОбъекта'] == disciplineCode:
            num = int(object['@Курс']) * 2 - 1 + int(object['@Семестр']) - 1
            if semesterNumberList.__contains__(num) == False:
                semesterNumberList.append(num)
    for i in range(semesterNumberList.__len__()):
        buffer = __SearchHoursBySemesterNumberOchnoe(semesterNumberList[i], disciplineCode, plxData)
        if bool(buffer):
            dict[semesterNumberList[i]] = buffer
    return dict


def __SearchHoursBySessionNumberZaochnoe(sessionNumber: int, disciplineCode: str, plxData: dict) -> dict:
    codeList = []
    hoursList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if (object['@КодТипаЧасов'] != '3' and object['@КодТипаЧасов'] != '5') and (object['@КодОбъекта'] == disciplineCode) and (
                int(object['@Курс']) * 2 - 1 + ((int(object['@Сессия']) - 1) // 2) == sessionNumber):
            if codeList.__contains__(object['@КодВидаРаботы']) == False:
                codeList.append(object['@КодВидаРаботы'])
                hoursList.append(object['@Количество'])
    nameList = []
    dict = {}
    for key in codeList:
        for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['СправочникВидыРабот']:
            if object['@Код'] == key:
                nameList.append(object['@Название'])
    for i in range(nameList.__len__()):
        dict[nameList[i]] = hoursList[i]

    return dict


def __SearchHoursZaochnoe(disciplineCode: str, plxData: dict) -> dict:
    dict = {}
    sessionNumbers = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if object['@КодОбъекта'] == disciplineCode and object['@Семестр'] == '0':
            num = int(object['@Курс']) * 2 - 1 + ((int(object['@Сессия']) - 1) // 2)
            if sessionNumbers.__contains__(num) == False:
                sessionNumbers.append(num)
    for key in sessionNumbers:
        buffer = __SearchHoursBySessionNumberZaochnoe(key, disciplineCode, plxData)
        if bool(buffer):
            dict[key] = buffer
    for key in dict.keys():
        dict[key]['ЗЕТ'] = str(int(dict[key]['Итого часов']) // 36)
    return dict

def __SearchPracticalJobsOchnoe (disciplineCode: str, plxData: dict) -> dict:
    dict = {}
    semesterNumberList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if object['@КодОбъекта'] == disciplineCode:
            num = int(object['@Курс']) * 2 - 1 + int(object['@Семестр']) - 1
            if semesterNumberList.__contains__(num) == False:
                semesterNumberList.append(num)
    for i in range(semesterNumberList.__len__()):
        buffer = __SearchPracticalJobsBySemesterNumber(semesterNumberList[i], disciplineCode, plxData)
        if bool(buffer):
            dict[semesterNumberList[i]] = buffer
    return dict


def __SearchPracticalJobsBySemesterNumber (semesterNumber: int, disciplineCode: str, plxData: dict) -> dict:
    codeList = []
    hoursList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if (object['@КодТипаЧасов'] == '5') and (
                object['@КодОбъекта'] == disciplineCode) and (
                int(object['@Курс']) * 2 - 1 + int(object['@Семестр']) - 1 == semesterNumber):
            if codeList.__contains__(object['@КодВидаРаботы']) == False:
                codeList.append(object['@КодВидаРаботы'])
                hoursList.append(object['@Количество'])
    nameList = []
    dict = {}
    for key in codeList:
        for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['СправочникВидыРабот']:
            if object['@Код'] == key:
                nameList.append(object['@Название'])
    for i in range(nameList.__len__()):
        dict[nameList[i]] = hoursList[i]

    return dict


def __SearchPracticalJobsZaochnoe (disciplineCode: str, plxData: dict) -> dict:
    dict = {}
    sessionNumber = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if object['@КодОбъекта'] == disciplineCode and object['@Семестр'] == '0':
            num = int(object['@Курс']) * 2 - 1 + ((int(object['@Сессия']) - 1) // 2)
            if sessionNumber.__contains__(num) == False:
                sessionNumber.append(num)
    for i in range(sessionNumber.__len__()):
        buffer = __SearchPracticalJobsBySessionNumber(sessionNumber[i], disciplineCode, plxData)
        if bool(buffer):
            dict[sessionNumber[i]] = buffer
    return dict


def __SearchPracticalJobsBySessionNumber (sessionNumber: int, disciplineCode: str, plxData: dict) -> dict:
    codeList = []
    hoursList = []
    for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['ПланыНовыеЧасы']:
        if (object['@КодТипаЧасов'] == '5') and (
                object['@КодОбъекта'] == disciplineCode) and (
                int(object['@Курс']) * 2 - 1 + ((int(object['@Сессия']) - 1) // 2) == sessionNumber):
            if codeList.__contains__(object['@КодВидаРаботы']) == False:
                codeList.append(object['@КодВидаРаботы'])
                hoursList.append(object['@Количество'])
    nameList = []
    dict = {}
    for key in codeList:
        for object in plxData['Документ']['diffgr:diffgram']['dsMMISDB']['СправочникВидыРабот']:
            if object['@Код'] == key:
                nameList.append(object['@Название'])
    for i in range(nameList.__len__()):
        dict[nameList[i]] = hoursList[i]
    return dict